# app/services/document_service.py
#
# Parses the ```generate-document fenced block the model emits (per the
# "Document Generation" section of SYSTEM_PROMPT in ai.py) and generates
# an actual PDF/DOCX/XLSX file from it. Called from chat.py once a full
# streamed response has been detected as a document block.

import re
import io
import logging

from fpdf import FPDF
from docx import Document
from openpyxl import Workbook

logger = logging.getLogger(__name__)

DOCUMENT_BLOCK_RE = re.compile(
    r"```generate-document\s*\n"
    r"format:\s*(?P<format>\w+)\s*\n"
    r"filename:\s*(?P<filename>[^\n]+)\s*\n"
    r"content:\s*\n"
    r"(?P<content>.*?)"
    r"```",
    re.DOTALL,
)

VALID_FORMATS = {"pdf", "docx", "xlsx"}


class DocumentParseError(Exception):
    pass


def parse_document_block(text: str) -> dict:
    """Extracts format/filename/content from a ```generate-document block.
    Raises DocumentParseError if the block is missing or malformed -
    callers should treat that as "not actually a valid document response"
    and fall back to showing the raw text."""
    match = DOCUMENT_BLOCK_RE.search(text)
    if not match:
        raise DocumentParseError("No generate-document block found")

    fmt = match.group("format").strip().lower()
    filename = match.group("filename").strip()
    content = match.group("content").strip("\n")

    if fmt not in VALID_FORMATS:
        raise DocumentParseError(f"Invalid format: {fmt}")
    if not filename:
        raise DocumentParseError("Missing filename")

    # Sanitize filename: strip anything that isn't alnum/dash/underscore,
    # since this ends up in a Content-Disposition header and (eventually)
    # a downloaded file name.
    safe_filename = re.sub(r"[^A-Za-z0-9\-_]", "-", filename)[:80] or "document"

    return {"format": fmt, "filename": safe_filename, "content": content}


def _generate_pdf(content: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)

    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 10, line[2:])
            pdf.set_font("Helvetica", size=11)
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, line[3:])
            pdf.set_font("Helvetica", size=11)
        elif line == "":
            pdf.ln(4)
        else:
            pdf.multi_cell(0, 6, line)

    return bytes(pdf.output())


def _generate_docx(content: str) -> bytes:
    doc = Document()
    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line == "":
            continue
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_xlsx(content: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    for raw_line in content.split("\n"):
        line = raw_line.rstrip()
        if not line:
            continue
        row = [cell.strip() for cell in line.split(",")]
        ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def generate_document(fmt: str, content: str) -> bytes:
    if fmt == "pdf":
        return _generate_pdf(content)
    elif fmt == "docx":
        return _generate_docx(content)
    elif fmt == "xlsx":
        return _generate_xlsx(content)
    raise DocumentParseError(f"Unsupported format: {fmt}")
